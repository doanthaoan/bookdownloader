"""
Updated version of wikicv_docx.py using SQLite database instead of HTML files for chapter lists
"""

import os
import re
import time
import random
import signal
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from colorama import Fore, init
from app.config import TRUYENWIKI, get_cookies, get_user_agent
from app.database import get_database
from app.services.text_cleaner import TextCleaner

# Initialize colorama
init(autoreset=True)

# Track active downloads for cancellation support
_active_downloads = {}

def register_download(book_id: int, downloader):
    _active_downloads[book_id] = downloader

def unregister_download(book_id: int):
    _active_downloads.pop(book_id, None)

def cancel_download(book_id: int):
    """Request cancellation of an active download."""
    downloader = _active_downloads.get(book_id)
    if downloader:
        downloader._cancel_flag = True
        return True
    return False

def cancel_all_downloads():
    """Cancel all active downloads (used on server shutdown)."""
    for book_id in list(_active_downloads.keys()):
        cancel_download(book_id)

def get_download_progress(book_id: int) -> dict:
    """Return progress info for an active download, or None if not active."""
    dl = _active_downloads.get(book_id)
    if not dl:
        return None
    return {
        "book_id": book_id,
        "cancelled": dl._cancel_flag,
        "success_count": getattr(dl, '_success_count', 0),
        "fail_count": getattr(dl, '_fail_count', 0),
        "total": len(getattr(dl, 'chapters', [])),
        "current_index": getattr(dl, '_current_index', 0),
        "current_title": getattr(dl, '_current_title', ''),
    }

class TruyenWikiDownloader:
    def __init__(self, book_name: str, redownload: bool = False):
        """
        Initialize downloader by book name instead of HTML file
        
        Args:
            book_name: Title of the book to download
            redownload: If True, use _redownload.docx output for failed chapters
        """
        self._cancel_flag = False
        self._redownload = redownload

        self.db = get_database()
        self.cleaner = TextCleaner()
        
        # Get book info from database
        self.book = self.db.get_book_by_title(book_name)
        if not self.book:
            raise ValueError(f"Book '{book_name}' not found in database. "
                           "Please run truyen_taodanhsach_db.py first to import books.")
        
        self.book_id = self.book['id']
        self.book_name = self.book['seo_title_basic']  # Use base_title for cleaner filenames
        self.domain = TRUYENWIKI['book_domain']
        
        # Load paths and settings from DB (with fallback to TRUYENWIKI config)
        self.book_path = self.db.get_setting('book_path') or TRUYENWIKI['book_path']
        self.logs_path = self.db.get_setting('logs_path') or TRUYENWIKI['logs_path']
        self.save_interval = int(self.db.get_setting('save_interval') or '10')
        self.page_load_timeout = int(self.db.get_setting('page_load_timeout') or '15')
        self.delay_min = int(self.db.get_setting('download_delay_min') or '4')
        self.delay_max = int(self.db.get_setting('download_delay_max') or '6')
        
        # Setup output paths
        if redownload:
            self.output_docx = os.path.join(self.book_path, f"{self.book_id}_{self.book_name}_redownload.docx")
        else:
            self.output_docx = os.path.join(self.book_path, f"{self.book_id}_{self.book_name}.docx")
        self.success_log = os.path.join(self.logs_path, f"{self.book_id}_{self.book_name}_success.txt")
        self.failure_log = os.path.join(self.logs_path, f"{self.book_id}_{self.book_name}_failure.txt")
        
        self._ensure_dirs()
        
        # Setup Selenium Driver
        self.driver = self._setup_selenium()
        self._inject_cookies()
        
        # Load existing DOCX if it exists to resume progress
        self.docx_doc = self._load_existing_docx()
        
        # Get chapters from database instead of HTML file
        self.chapters = self.db.get_chapters_by_book(self.book_id)
        if not self.chapters:
            print(f"{Fore.YELLOW}Warning: No chapters found for book '{book_name}'")
            print(f"{Fore.YELLOW}You may need to run chapter extraction first.")
    
    def _ensure_dirs(self):
        """Ensure required directories exist"""
        for path in [self.book_path, self.logs_path]:
            os.makedirs(path, exist_ok=True)
    
    def _load_existing_docx(self):
        """Load existing DOCX file if it exists, or create a fresh one with the book title."""
        fresh = False
        if os.path.exists(self.output_docx):
            print(f"{Fore.CYAN}Found existing DOCX file. Resuming progress...")
            doc = Document(self.output_docx)
        else:
            doc = Document()
            fresh = True

        if fresh:
            # Remove the default empty paragraph and add book title
            for p in list(doc.paragraphs):
                p._element.getparent().remove(p._element)
            doc.add_heading(self.book['title'], level=0)

        return doc
    
    def _setup_selenium(self):
        """Setup Selenium WebDriver with appropriate options"""
        chrome_options = Options()
        # chrome_options.add_argument("--headless")  # Run in background
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-blink-features=AutomationControlled") 
        chrome_options.add_argument(f"user-agent={get_user_agent()}")
        
        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()),
            options=chrome_options
        )
        return driver
    
    def _inject_cookies(self):
        """Inject cookies into the browser session."""
        print(f"{Fore.CYAN}Injecting cookies...")
        self.driver.get(self.domain)
        
        cookie_map = get_cookies()
        print(f"{Fore.CYAN}Using cookies: {cookie_map}")
        
        cookie_domain = TRUYENWIKI['cookie_domain']
        for name, value in cookie_map.items():
            if value:
                try:
                    self.driver.add_cookie({'name': name, 'value': value, 'domain': cookie_domain})
                except Exception as e:
                    print(f"{Fore.YELLOW}Could not add cookie {name}: {e}")

    def log_event(self, log_path: str, message: str):
        """Log an event to the specified log file (write-only audit trail)."""
        with open(log_path, 'a', encoding='utf-8') as f:
            f.write(f"{time.strftime('%X %x')}: {message}\n")

    def process_chapter(self, chapter: dict) -> tuple:
        """
        Process a single chapter
        
        Args:
            chapter: Dictionary with chapter info from database
            
        Returns:
            Tuple of (chapter_title, full_url)
        """
        # Build full URL
        chapter_url = chapter['chapter_url']
        full_url = chapter_url if chapter_url.startswith("http") else self.domain + chapter_url
        
        # Navigate to the chapter
        self.driver.get(full_url)
        
        # WAIT: Wait for the content wrapper to appear (timeout from DB settings)
        wait = WebDriverWait(self.driver, self.page_load_timeout)
        wait.until(EC.presence_of_element_located((By.CLASS_NAME, "content-body-wrapper")))
        
        # Anti-bot cooldown: random delay from DB settings
        time.sleep(random.randint(self.delay_min, self.delay_max))
        
        # Get the rendered HTML
        soup = BeautifulSoup(self.driver.page_source, 'html.parser')
        
        # Extract Title and Content
        title_tags = soup.find_all('p', {'class': 'book-title'})
        title_tag = title_tags[1] if len(title_tags) > 1 else title_tags[0]
        content_tag = soup.find('div', {'class': 'content-body-wrapper'})
        
        if not content_tag:
            raise Exception("Could not find content-body-wrapper after waiting.")

        title_text = self.cleaner.clean(title_tag.get_text())
        
        # Save to HTML (Appended mode)
        # with open(self.output_html, 'a', encoding='utf-8') as html_file:
        #     if os.path.getsize(self.output_html) == 0:
        #         html_file.write('<html><body>\n')
        #     html_file.write(f"<h1>{title_text}</h1>\n")
        #     html_file.write(str(content_tag))
        
        # Add to DOCX object
        self.docx_doc.add_heading(title_text, level=1)
        for p in content_tag.find_all('p'):
            self.docx_doc.add_paragraph(self.cleaner.clean(p.get_text()))

        return title_text, full_url

    def run(self, max_chapters=None):
        """Main download process.
        
        Args:
            max_chapters: Maximum number of chapters to process in this session.
                          None = no limit.
        """
        if not self.chapters:
            print(f"{Fore.RED}No chapters to process for book '{self.book_name}'")
            print(f"{Fore.YELLOW}Please extract chapters first using the chapter extractor.")
            self.driver.quit()
            return

        # Register for cancellation support
        register_download(self.book_id, self)

        # Mark as in-progress in DB
        self.db.update_book_status(
            book_id=self.book_id,
            download_status='in_progress',
        )

        # Filter to only pending chapters (DB is the source of truth)
        chapters_to_process = [c for c in self.chapters if c['download_status'] == 'pending']

        if not chapters_to_process:
            print(f"{Fore.GREEN}All chapters already downloaded for '{self.book_name}'!")
            self.driver.quit()
            unregister_download(self.book_id)
            return

        completed = len(self.chapters) - len(chapters_to_process)
        print(f"📚 Found {len(self.chapters)} total chapters. "
              f"{completed} already done. "
              f"{len(chapters_to_process)} remaining.")

        limited = False
        if max_chapters and len(chapters_to_process) > max_chapters:
            chapters_to_process = chapters_to_process[:max_chapters]
            limited = True
            print(f"🔢 Session limited to {max_chapters} chapters.")
        
        self.log_event(self.success_log, "RESUMING/STARTING SESSION")
        
        success_count = 0
        fail_count = 0
        start_total = time.time()
        cancelled = False

        # Expose progress for real-time tracking
        self._success_count = 0
        self._fail_count = 0
        self._current_index = 0
        self._current_title = ''

        try:
            for index, chapter in enumerate(chapters_to_process, 1):
                self._current_index = index
                self._current_title = chapter['chapter_title']

                # Check for cancellation
                if self._cancel_flag:
                    print(f"{Fore.YELLOW}Download cancelled by user.")
                    cancelled = True
                    break

                try:
                    title, full_url = self.process_chapter(chapter)
                    
                    # Log success
                    self.log_event(self.success_log, full_url)
                    
                    # Update chapter status in database
                    self.db.update_chapter_status(
                        chapter_id=chapter['id'],
                        status='completed',
                        file_path=f"{self.book_name}.docx"
                    )
                    
                    success_count += 1
                    self._success_count = success_count
                    
                    # Save checkpoint every save_interval chapters
                    if self.save_interval and success_count % self.save_interval == 0:
                        self.docx_doc.save(self.output_docx)
                        print(f"{Fore.CYAN}--- Checkpoint: DOCX saved automatically ---")
                    
                    print(f"{Fore.GREEN}Success: {Fore.WHITE}{title} "
                          f"({index}/{len(chapters_to_process)})")
                    
                except Exception as e:
                    # Log failure
                    error_msg = f"{chapter['chapter_url']} | Error: {e}"
                    self.log_event(self.failure_log, error_msg)
                    
                    # Update chapter status in database
                    self.db.update_chapter_status(
                        chapter_id=chapter['id'],
                        status='failed',
                        error_message=str(e)
                    )
                    
                    fail_count += 1
                    self._fail_count = fail_count
                    print(f"{Fore.RED}Failed: {chapter['chapter_title']} | Error: {e}")
                    
                    # Save DOCX after failure to preserve progress
                    self.docx_doc.save(self.output_docx)
                
                # Add delay between chapters to be respectful to the server
                if index < len(chapters_to_process):  # No delay after last chapter
                    delay = random.randint(2, 4)
                    time.sleep(delay)

        finally:
            # Save docx before anything else (critical on interrupt/cancel)
            self.docx_doc.save(self.output_docx)

            # Always close the browser window
            self.driver.quit()

        total_duration = time.time() - start_total
        
        # Update book status in database
        if cancelled:
            new_status = 'cancelled'
        elif limited:
            new_status = 'paused'
        elif fail_count == 0:
            new_status = 'completed'
        elif success_count == 0:
            new_status = 'failed'
        else:
            new_status = 'completed_with_errors'
            
        self.db.update_book_status(
            book_id=self.book_id,
            download_status=new_status,
            downloaded_chapters=success_count
        )
        
        self._print_summary(success_count, fail_count, len(chapters_to_process), total_duration)

        # Unregister after completion
        unregister_download(self.book_id)

    def run_redownload(self, all_chapters: bool = False):
        """
        Re-download chapters into a separate _redownload.docx.
        
        Args:
            all_chapters: If True, re-download ALL chapters (fresh copy).
                         If False, only failed chapters.
        """
        if not self.chapters:
            print(f"{Fore.RED}No chapters found for '{self.book_name}'")
            self.driver.quit()
            return

        register_download(self.book_id, self)
        self.db.update_book_status(book_id=self.book_id, download_status='in_progress')

        if all_chapters:
            chapters_to_process = list(self.chapters)
            label = "all"
        else:
            chapters_to_process = [c for c in self.chapters if c['download_status'] == 'failed']
            label = "failed"

        if not chapters_to_process:
            print(f"{Fore.GREEN}No {label} chapters to re-download.")
            self.driver.quit()
            unregister_download(self.book_id)
            return

        print(f"{Fore.YELLOW}Re-downloading {len(chapters_to_process)} {label} chapter(s) into {self.output_docx}")
        success_count = 0
        fail_count = 0
        start_total = time.time()
        cancelled = False

        self._success_count = 0
        self._fail_count = 0
        self._current_index = 0
        self._current_title = ''

        try:
            for index, chapter in enumerate(chapters_to_process, 1):
                self._current_index = index
                self._current_title = chapter['chapter_title']

                if self._cancel_flag:
                    print(f"{Fore.YELLOW}Re-download cancelled by user.")
                    cancelled = True
                    break

                try:
                    title, full_url = self.process_chapter(chapter)
                    self.log_event(self.success_log, f"[REDOWNLOAD] {full_url}")
                    self.db.update_chapter_status(
                        chapter_id=chapter['id'], status='completed',
                        file_path=f"{self.book_name}_redownload.docx"
                    )
                    success_count += 1
                    self._success_count = success_count

                    if self.save_interval and success_count % self.save_interval == 0:
                        self.docx_doc.save(self.output_docx)

                    print(f"{Fore.GREEN}Redownload success: {title} ({index}/{len(chapters_to_process)})")

                except Exception as e:
                    error_msg = f"{chapter['chapter_url']} | Error: {e}"
                    self.log_event(self.failure_log, error_msg)
                    self.db.update_chapter_status(
                        chapter_id=chapter['id'], status='failed', error_message=str(e)
                    )
                    fail_count += 1
                    self._fail_count = fail_count
                    print(f"{Fore.RED}Redownload failed: {chapter['chapter_title']} | Error: {e}")
                    self.docx_doc.save(self.output_docx)

                if index < len(chapters_to_process):
                    time.sleep(random.randint(2, 4))

        finally:
            self.docx_doc.save(self.output_docx)
            self.driver.quit()

        total_duration = time.time() - start_total
        new_status = 'cancelled' if cancelled else ('completed' if fail_count == 0 else 'completed_with_errors')
        self.db.update_book_status(
            book_id=self.book_id, download_status=new_status,
            downloaded_chapters=len([c for c in self.chapters if c['download_status'] == 'completed'])
        )
        self._print_summary(success_count, fail_count, len(chapters_to_process), total_duration)
        unregister_download(self.book_id)

    def run_single(self, chapter_id: int):
        """Download a single chapter and append to _redownload.docx."""
        chapter = next((c for c in self.chapters if c['id'] == chapter_id), None)
        if not chapter:
            print(f"{Fore.RED}Chapter id={chapter_id} not found for book '{self.book_name}'")
            self.driver.quit()
            return

        register_download(self.book_id, self)
        self._success_count = 0
        self._fail_count = 0
        self._current_index = 1
        self._current_title = chapter['chapter_title']

        try:
            title, full_url = self.process_chapter(chapter)
            self.log_event(self.success_log, f"[SINGLE] {full_url}")
            self.db.update_chapter_status(
                chapter_id=chapter['id'], status='completed',
                file_path=f"{self.book_name}_redownload.docx"
            )
            self._success_count = 1
            self.docx_doc.save(self.output_docx)
            print(f"{Fore.GREEN}Single download success: {title}")

        except Exception as e:
            error_msg = f"{chapter['chapter_url']} | Error: {e}"
            self.log_event(self.failure_log, error_msg)
            self.db.update_chapter_status(
                chapter_id=chapter['id'], status='failed', error_message=str(e)
            )
            self._fail_count = 1
            self.docx_doc.save(self.output_docx)
            print(f"{Fore.RED}Single download failed: {chapter['chapter_title']} | Error: {e}")
            raise

        finally:
            self.driver.quit()
            unregister_download(self.book_id)

    def _print_summary(self, success: int, fail: int, total: int, duration: float):
        print("\n" + "="*50)
        print("TỔNG KẾT PHIÊN TẢI")
        print("="*50)
        print(f"{Fore.GREEN}Thành công: {success}/{total}")
        print(f"{Fore.RED}Thất bại: {fail}/{total}")
        print(f"{Fore.CYAN}Thời gian chạy: {duration:.2f} s")
        print(f"{Fore.YELLOW}Sách: {self.book_name}")
        print(f"{Fore.BLUE}File DOCX: {self.output_docx}")
        # print(f"{Fore.BLUE}File HTML: {self.output_html}")

# Convenience function for easy usage
def download_book(book_name: str, max_chapters: int = None):
    """
    Download a book by name

    Args:
        book_name: Title of the book to download
        max_chapters: Maximum number of chapters to process (None = no limit)
    """
    downloader = TruyenWikiDownloader(book_name)
    try:
        downloader.run(max_chapters=max_chapters)
    except Exception as e:
        print(f"❌ Error downloading book '{book_name}': {e}")
        raise


def redownload_book(book_name: str, all_chapters: bool = False):
    """Re-download chapters for a book into _redownload.docx."""
    downloader = TruyenWikiDownloader(book_name, redownload=True)
    try:
        downloader.run_redownload(all_chapters=all_chapters)
    except Exception as e:
        print(f"❌ Error re-downloading book '{book_name}': {e}")
        raise


def download_single_chapter(book_name: str, chapter_id: int):
    """Download a single chapter and append to _redownload.docx."""
    downloader = TruyenWikiDownloader(book_name, redownload=True)
    try:
        downloader.run_single(chapter_id)
    except Exception as e:
        print(f"❌ Error downloading chapter {chapter_id} of '{book_name}': {e}")
        raise

if __name__ == "__main__":
    # Ask for book name
    book_name = input("Tên truyện: ")
    download_book(book_name)